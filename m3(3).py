import docx
new_document=docx.Document(r"C:\Users\U1092451\Desktop\个人文件\培训\SOP表格使用方法\空白模板\UZA77631_Curriculum Vitae of Bin Yang.docx")
#table[0].cell(1,0)是姓名
new_document.tables[0].cell(1,0).text='abc'
#table[0],cell(1,1)是执业证书
new_document.tables[0].cell(1,1).text='abc'
#table[1],cell(1,0)是英文地址
new_document.tables[1].cell(1,0).text='abc'
#table[2],cell(1,0)是电话
new_document.tables[2].cell(1,0).text='abc'
#table[2],cell(1,1)是传真，NA
new_document.tables[2].cell(1,1).text='NA'
#table[3],cell(1,0)是邮箱
new_document.tables[3].cell(1,0).text='abc'

#table[4],cell(2,0)是学习经历1的学校
new_document.tables[4].cell(2,0).text='abc'
#table[4],cell(2,1)是学习经历1的学位和时间
new_document.tables[4].cell(2,1).text='abc'
#table[4],cell(2,2)是学习经历1的学习领域
new_document.tables[4].cell(2,2).text='abc'
#table[4],cell(3,0)是学习经历2的学校
new_document.tables[4].cell(3,0).text='abc'
#table[4],cell(3,1)是学习经历2的学位和时间
new_document.tables[4].cell(3,1).text='abc'
#table[4],cell(3,2)是学习经历2的学习领域
new_document.tables[4].cell(3,2).text='abc'

#table[6],cell(2,0)是工作经历1的职位
current=new_document.tables[6].cell(1,0).text
new_document.tables[6].cell(1,0).text=current+'\n'+'abc'
#table[6],cell(2,1)是工作经历1的工作地点
new_document.tables[6].cell(1,1).text='abc'
#table[6],cell(2,2)是工作经历1的起始终止时间
new_document.tables[6].cell(1,2).text='abc'
#table[6],cell(2,0)是工作经历2的职位
previous=new_document.tables[6].cell(2,0).text
new_document.tables[6].cell(2,0).text=previous+'\n'+'abc'
#table[6],cell(2,1)是工作经历2的工作地点
new_document.tables[6].cell(2,1).text='abc'
#table[6],cell(2,2)是工作经历2的起始终止时间
new_document.tables[6].cell(2,2).text='abc'

#table[7],cell(2,0)是临床研究经历的领域
new_document.tables[7].cell(2,0).text='abc'
#table[7],cell(2,1)是临床研究经历的分期
new_document.tables[7].cell(2,1).text='abc'
#table[7],cell(2,2)是临床研究经历的角色
new_document.tables[7].cell(2,2).text='abc'
#table[7],cell(2,3)是临床研究经历的时间
new_document.tables[7].cell(2,3).text='abc'
#table[7],cell(3,0)是临床研究经历的领域
new_document.tables[7].cell(3,0).text='abc'
#table[7],cell(3,1)是临床研究经历的分期
new_document.tables[7].cell(3,1).text='abc'
#table[7],cell(3,2)是临床研究经历的角色
new_document.tables[7].cell(3,2).text='abc'
#table[7],cell(3,3)是临床研究经历的时间
new_document.tables[7].cell(3,3).text='abc'

#table[8],cell(1,0)是ICH-GCP的培训情况
ICH_GCP=new_document.tables[8].cell(1,0).text
new_document.tables[8].cell(1,0).text=ICH_GCP+'\n'+'abc'

#table[9],cell(1,0)是其它经历的培训情况
new_document.tables[9].cell(1,0).text='abc'

#保存文件
new_document.save(r"C:\Users\U1092451\Desktop\个人文件\培训\SOP表格使用方法\空白模板\UZA77631_Curriculum Vitae of Bin Yang.docx")